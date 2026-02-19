import zipfile
import xml.etree.ElementTree as ET
from docx import Document
import io
import logging
from typing import List

logger = logging.getLogger(__name__)

class DOCXParser:
    def __init__(self):
        self.ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    def parse(self, file_bytes: bytes):
        """
        Parses a DOCX file and analyzes it for ATS risks.
        """
        stream = io.BytesIO(file_bytes)
        
        # 1. Standard Text Extraction
        try:
            doc = Document(stream)
            full_text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            extracted_text = '\n'.join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            extracted_text = ""
            return {"error": "Failed to parse DOCX"}

        # 2. XML Analysis for Floating Objects & Headers
        # We need to unzip and read word/document.xml directly to find <w:textBox>
        stream.seek(0)
        floating_object_count = 0
        has_header_footer_content = False
        extracted_urls: List[str] = []
        
        try:
            with zipfile.ZipFile(stream) as zf:
                xml_content = zf.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Count text boxes
                # w:textBox is often inside w:pict or w:drawing
                text_boxes = tree.findall('.//w:textBox', self.ns)
                floating_object_count = len(text_boxes)
                
                # Check for header/footer references in the document relation
                # This is a proxy; if there are refs, there might be content.
                # A better check is to see if header files exist and have text.
                # But for now, we just check if the XML has headerReference tags.
                header_refs = tree.findall('.//w:headerReference', self.ns)
                footer_refs = tree.findall('.//w:footerReference', self.ns)
                
                if header_refs or footer_refs:
                    # Check if they actually contain text
                    has_header_footer_content = self._check_header_footer_content(zf)

                extracted_urls = self._extract_hyperlinks(zf)

        except Exception as e:
            logger.error(f"XML analysis failed: {e}")

        return {
            "raw_text": extracted_text,
            "floating_object_count": floating_object_count,
            "has_header_footer_content": has_header_footer_content,
            "file_size_bytes": len(file_bytes),
            "extracted_urls": extracted_urls,
        }

    def _check_header_footer_content(self, zf):
        """
        Checks if any header or footer xml files in the zip actually contain text.
        """
        try:
            for name in zf.namelist():
                if name.startswith('word/header') or name.startswith('word/footer'):
                    xml = zf.read(name)
                    if b'<w:t>' in xml or b'<w:t ' in xml: # Quick check for text tags
                        return True
        except:
            pass
        return False

    def _extract_hyperlinks(self, zf) -> List[str]:
        """Extract external hyperlinks from DOCX relationship files."""
        rel_ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        urls: List[str] = []
        seen = set()

        for name in zf.namelist():
            if not name.endswith(".rels"):
                continue
            try:
                xml = zf.read(name)
                root = ET.fromstring(xml)
            except Exception:
                continue

            for rel in root.findall('r:Relationship', rel_ns):
                target = str(rel.attrib.get("Target", "")).strip()
                target_mode = str(rel.attrib.get("TargetMode", "")).strip().lower()
                if not target or target_mode != "external":
                    continue
                if not target.startswith(("http://", "https://", "www.", "linkedin.com", "github.com")):
                    continue
                cleaned = target.rstrip(".,);")
                if cleaned.startswith("www.") or cleaned.startswith("linkedin.com") or cleaned.startswith("github.com"):
                    cleaned = f"https://{cleaned}"
                if cleaned in seen:
                    continue
                seen.add(cleaned)
                urls.append(cleaned)

        return urls[:40]
