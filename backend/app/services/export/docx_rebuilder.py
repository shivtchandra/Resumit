"""
DOCX Rebuilder
Reconstructs DOCX files from layout schemas while preserving formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)


class DOCXRebuilder:
    def __init__(self):
        """Initialize DOCX rebuilder."""
        self.default_font = "Calibri"
        self.default_font_size = 11
        self.heading_font_size = 14
        self.name_font_size = 16
    
    def _sanitize_text(self, text: str) -> str:
        """
        Sanitize text to remove control characters and NULL bytes.
        Keeps only valid XML characters.
        """
        if not text:
            return ""
        
        # Remove NULL bytes and control characters except tab, newline, carriage return
        sanitized = ''.join(
            char for char in text
            if ord(char) >= 32 or char in '\t\n\r'
        )
        
        return sanitized
    
    def rebuild_from_schema(self, schema: Dict[str, Any], output_path: str) -> str:
        """
        Rebuild DOCX file from layout schema.
        
        Args:
            schema: Layout schema with sections
            output_path: Path to save DOCX file
            
        Returns:
            Path to created DOCX file
        """
        doc = Document()
        
        # Set default style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.default_font
        font.size = Pt(self.default_font_size)
        
        # Process sections in order
        sections = sorted(schema.get("sections", []), key=lambda s: s.get("pos", 0))
        
        for section in sections:
            section_type = section.get("type")
            
            if section_type == "CONTACT":
                self._add_contact_section(doc, section)
            elif section_type == "SUMMARY":
                self._add_summary_section(doc, section)
            elif section_type == "EXPERIENCE":
                self._add_experience_section(doc, section)
            elif section_type == "EDUCATION":
                self._add_education_section(doc, section)
            elif section_type == "SKILLS":
                self._add_skills_section(doc, section)
            elif section_type == "PROJECTS":
                self._add_projects_section(doc, section)
            elif section_type == "CERTIFICATIONS":
                self._add_certifications_section(doc, section)
            else:
                # Generic section
                self._add_generic_section(doc, section)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Saved DOCX to {output_path}")
        
        return output_path
    
    def _add_contact_section(self, doc: Document, section: Dict[str, Any]):
        """Add contact information section."""
        raw_text = section.get("raw", "")
        parsed = section.get("parsed", {})
        
        # Add name (first line, usually)
        lines = raw_text.split('\n')
        if lines:
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(lines[0].strip())
            name_run.font.size = Pt(self.name_font_size)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact details
        contact_parts = []
        if parsed.get("email"):
            contact_parts.append(parsed["email"])
        if parsed.get("phone"):
            contact_parts.append(parsed["phone"])
        if parsed.get("linkedin"):
            contact_parts.append(parsed["linkedin"])
        if parsed.get("github"):
            contact_parts.append(parsed["github"])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_summary_section(self, doc: Document, section: Dict[str, Any]):
        """Add professional summary section."""
        # Section header
        heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Summary content
        summary_text = self._sanitize_text(section.get("raw", ""))
        if summary_text:
            para = doc.add_paragraph(summary_text)
            para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_experience_section(self, doc: Document, section: Dict[str, Any]):
        """Add work experience section."""
        # Section header
        heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Add each experience entry
        for entry in section.get("entries", []):
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{entry.get('title', 'Unknown Title')}")
            title_run.font.bold = True
            title_run.font.size = Pt(self.default_font_size)
            
            # Company and dates
            company_para = doc.add_paragraph()
            company_text = f"{entry.get('company', 'Unknown Company')}"
            dates_text = f"{entry.get('start', '')} - {entry.get('end', 'Present')}"
            
            company_run = company_para.add_run(company_text)
            company_run.font.italic = True
            company_run.font.size = Pt(self.default_font_size)
            
            company_para.add_run(" | ")
            
            dates_run = company_para.add_run(dates_text)
            dates_run.font.size = Pt(self.default_font_size)
            
            # Bullets
            for bullet in entry.get("bullets", []):
                clean_bullet = self._sanitize_text(bullet)
                if clean_bullet:
                    bullet_para = doc.add_paragraph(clean_bullet, style='List Bullet')
                    bullet_para.runs[0].font.size = Pt(self.default_font_size)
            
            # Spacing between entries
            doc.add_paragraph()
    
    def _add_education_section(self, doc: Document, section: Dict[str, Any]):
        """Add education section."""
        # Section header
        heading = doc.add_heading('EDUCATION', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Check if we have parsed entries or raw text
        if section.get("entries"):
            for entry in section["entries"]:
                # Degree
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(entry.get("degree", ""))
                degree_run.font.bold = True
                degree_run.font.size = Pt(self.default_font_size)
                
                # Institution and year
                inst_para = doc.add_paragraph()
                inst_text = entry.get("institution", "")
                year_text = entry.get("year", "")
                
                if inst_text:
                    inst_run = inst_para.add_run(inst_text)
                    inst_run.font.italic = True
                    inst_run.font.size = Pt(self.default_font_size)
                
                if year_text:
                    if inst_text:
                        inst_para.add_run(" | ")
                    year_run = inst_para.add_run(year_text)
                    year_run.font.size = Pt(self.default_font_size)
                
                # Details
                for detail in entry.get("details", []):
                    detail_para = doc.add_paragraph(detail, style='List Bullet')
                    detail_para.runs[0].font.size = Pt(self.default_font_size)
        else:
            # Raw text
            raw_text = section.get("raw", "")
            para = doc.add_paragraph(raw_text)
            para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_skills_section(self, doc: Document, section: Dict[str, Any]):
        """Add skills section."""
        # Section header
        heading = doc.add_heading('SKILLS', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Skills content
        skills_text = self._sanitize_text(section.get("raw", ""))
        if skills_text:
            para = doc.add_paragraph(skills_text)
            para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_projects_section(self, doc: Document, section: Dict[str, Any]):
        """Add projects section."""
        # Section header
        heading = doc.add_heading('PROJECTS', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Projects content
        projects_text = self._sanitize_text(section.get("raw", ""))
        if projects_text:
            para = doc.add_paragraph(projects_text)
            para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_certifications_section(self, doc: Document, section: Dict[str, Any]):
        """Add certifications section."""
        # Section header
        heading = doc.add_heading('CERTIFICATIONS', level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Certifications content
        certs_text = section.get("raw", "")
        para = doc.add_paragraph(certs_text)
        para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_generic_section(self, doc: Document, section: Dict[str, Any]):
        """Add generic section."""
        section_type = section.get("type", "UNKNOWN")
        
        # Section header
        heading = doc.add_heading(section_type.upper(), level=2)
        heading.runs[0].font.size = Pt(self.heading_font_size)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Content
        raw_text = section.get("raw", "")
        para = doc.add_paragraph(raw_text)
        para.runs[0].font.size = Pt(self.default_font_size)
        
        # Add spacing
        doc.add_paragraph()
